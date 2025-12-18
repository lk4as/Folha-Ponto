{\rtf1\ansi\ansicpg1252\cocoartf2865
\cocoatextscaling0\cocoaplatform0{\fonttbl\f0\fswiss\fcharset0 Helvetica;}
{\colortbl;\red255\green255\blue255;}
{\*\expandedcolortbl;;}
\paperw11900\paperh16840\margl1440\margr1440\vieww11520\viewh8400\viewkind0
\pard\tx720\tx1440\tx2160\tx2880\tx3600\tx4320\tx5040\tx5760\tx6480\tx7200\tx7920\tx8640\pardirnatural\partightenfactor0

\f0\fs24 \cf0 import streamlit as st\
from docx import Document\
from docx.shared import Pt, Inches, Cm\
from docx.oxml.ns import qn\
from datetime import date, datetime, timedelta\
import io\
\
# --- CONFIGURA\'c7\'d5ES VISUAIS ---\
st.set_page_config(page_title="Gerador de Folha de Ponto", page_icon="\uc0\u55357 \u56541 ")\
\
st.title("\uc0\u55357 \u56541  Gerador de Folha de Ponto - Bram Offshore")\
st.markdown("Preencha os dados abaixo para gerar o documento automaticamente.")\
\
# --- ARQUIVOS (Est\'e3o na mesma pasta do script) ---\
ARQUIVO_MODELO = "Folha de Ponto.docx"\
ARQUIVO_ASSINATURA = "assinatura.png"\
NOME_DA_FONTE = 'Arial'\
\
# S\'edmbolos\
CHECKBOX_MARCADO = "\\u2612"\
CHECKBOX_VAZIO = "\\u2610"\
\
# --- FUN\'c7\'d5ES UTILIT\'c1RIAS (A mesma l\'f3gica robusta) ---\
def configurar_fonte(run, tamanho=9):\
    run.font.name = NOME_DA_FONTE\
    run.font.size = Pt(tamanho)\
    r = run._element\
    r.rPr.rFonts.set(qn('w:eastAsia'), NOME_DA_FONTE)\
\
def limpar_e_escrever(celula, texto, tamanho=9, alinhamento=1):\
    celula._element.clear_content()\
    p = celula.add_paragraph()\
    p.alignment = alinhamento\
    run = p.add_run(str(texto))\
    configurar_fonte(run, tamanho)\
\
def inserir_assinatura(celula, caminho_imagem):\
    celula._element.clear_content()\
    p = celula.add_paragraph()\
    p.alignment = 1\
    run = p.add_run()\
    try:\
        run.add_picture(caminho_imagem, width=Inches(0.9))\
    except FileNotFoundError:\
         run.add_text("[Assinatura]")\
         configurar_fonte(run, 8)\
\
def iterar_todas_as_tabelas(doc_obj):\
    for table in doc_obj.tables:\
        yield table\
        for row in table.rows:\
            for cell in row.cells:\
                if cell.tables:\
                    yield from iterar_todas_as_tabelas(cell)\
\
def preencher_campo_seguro(doc, rotulo_busca, valor_preencher, tamanho_fonte=10):\
    rotulo_limpo = rotulo_busca.lower()\
    for table in iterar_todas_as_tabelas(doc):\
        for row in table.rows:\
            for i, cell in enumerate(row.cells):\
                if rotulo_limpo in cell.text.lower():\
                    for j in range(i + 1, len(row.cells)):\
                        cand_cell = row.cells[j]\
                        if cand_cell != cell and rotulo_limpo not in cand_cell.text.lower():\
                            limpar_e_escrever(cand_cell, valor_preencher, tamanho=tamanho_fonte, alinhamento=0)\
                            return True\
    return False\
\
def marcar_base_preservando_linhas(doc, texto_base_alvo):\
    for table in iterar_todas_as_tabelas(doc):\
        for row in table.rows:\
            for cell in row.cells:\
                if "(Rio)" in cell.text and "CNPJ" not in cell.text: \
                    linhas_originais = [p.text for p in cell.paragraphs if p.text.strip()]\
                    cell._element.clear_content()\
                    for linha in linhas_originais:\
                        p = cell.add_paragraph()\
                        p.alignment = 0 \
                        p.paragraph_format.space_after = Pt(0)\
                        texto_limpo = linha.replace(CHECKBOX_MARCADO, "").replace(CHECKBOX_VAZIO, "").replace("\uc0\u9746 ", "").replace("\u9744 ", "").strip()\
                        if texto_base_alvo in linha:\
                            run = p.add_run(f"\{CHECKBOX_MARCADO\} \{texto_limpo\}")\
                            configurar_fonte(run, 8)\
                        else:\
                            run = p.add_run(f"\{CHECKBOX_VAZIO\} \{texto_limpo\}")\
                            configurar_fonte(run, 8)\
                    return True\
    return False\
\
def sanitarizar_hora(hora_input):\
    if not hora_input: return "00:00"\
    h = str(hora_input).strip()\
    if ":" not in h:\
        try:\
            return f"\{int(h):02d\}:00"\
        except ValueError:\
            return h\
    return h\
\
def calcular_quarteto(hora_entrada, hora_almoco_ida):\
    fmt = "%H:%M"\
    hora_entrada = sanitarizar_hora(hora_entrada)\
    hora_almoco_ida = sanitarizar_hora(hora_almoco_ida)\
    \
    try:\
        dt_ent = datetime.strptime(hora_entrada, fmt)\
        dt_alm_ida = datetime.strptime(hora_almoco_ida, fmt)\
        dt_alm_volta = dt_alm_ida + timedelta(hours=1)\
        dt_saida = dt_ent + timedelta(hours=7)\
        return [hora_entrada, hora_almoco_ida, dt_alm_volta.strftime(fmt), dt_saida.strftime(fmt)]\
    except ValueError:\
        return ["--:--", "--:--", "--:--", "--:--"]\
\
def forcar_uma_pagina(doc):\
    for section in doc.sections:\
        section.bottom_margin = Cm(0.5)\
        section.footer_distance = Cm(0.5)\
    if doc.paragraphs:\
        p = doc.paragraphs[-1]\
        p.paragraph_format.space_after = Pt(0)\
        for run in p.runs: run.font.size = Pt(1)\
\
# --- INTERFACE DO STREAMLIT ---\
\
col1, col2 = st.columns(2)\
with col1:\
    mes_final = st.selectbox("M\'eas", options=list(range(1, 13)), index=11, format_func=lambda x: f"\{x:02d\}")\
with col2:\
    ano_final = st.number_input("Ano", value=2025, step=1)\
\
col3, col4 = st.columns(2)\
with col3:\
    nome_func = st.text_input("Nome do Funcion\'e1rio", "Lukas Souza Henriques Crespo")\
with col4:\
    cargo_func = st.text_input("Fun\'e7\'e3o / Job Title", "Estagi\'e1rio - Dep. DP Assurance")\
\
col5, col6 = st.columns(2)\
with col5:\
    base_opt = st.selectbox("Base", ["Rio", "A\'e7u", "Maca\'e9", "Guaxindiba"], index=1)\
    mapa_bases = \{"Rio": "(Rio)", "A\'e7u": "(A\'e7u)", "Maca\'e9": "(Maca\'e9)", "Guaxindiba": "(Guax.)"\}\
    texto_base = mapa_bases[base_opt]\
with col6:\
    data_emissao = st.text_input("Data de Emiss\'e3o (CTPS)", "04/02/2020")\
\
st.markdown("---")\
st.subheader("\uc0\u9200  Configura\'e7\'e3o de Hor\'e1rios (Carga 7h)")\
\
tipo_horario = st.radio("Tipo de Hor\'e1rio", ["Fixo (Igual todos os dias)", "Vari\'e1vel (Muda na semana)"], horizontal=True)\
\
escala_semanal = \{\}\
txt_horario_cabecalho = ""\
\
if tipo_horario.startswith("Fixo"):\
    c1, c2 = st.columns(2)\
    ent = c1.text_input("Entrada (ex: 8 ou 08:00)", "08:00")\
    alm = c2.text_input("Sa\'edda Almo\'e7o (ex: 12 ou 12:00)", "12:00")\
    \
    quarteto = calcular_quarteto(ent, alm)\
    st.info(f"Escala calculada: \{quarteto[0]\} - \{quarteto[3]\} (Almo\'e7o: \{quarteto[1]\} \'e0s \{quarteto[2]\})")\
    \
    for d in range(5): escala_semanal[d] = quarteto\
    txt_horario_cabecalho = f"\{quarteto[0]\} - \{quarteto[3]\}"\
\
else:\
    st.write("Defina a entrada de cada dia e o hor\'e1rio de almo\'e7o padr\'e3o.")\
    cols = st.columns(5)\
    dias = ["Seg", "Ter", "Qua", "Qui", "Sex"]\
    entradas = []\
    \
    for i, dia in enumerate(dias):\
        val = cols[i].text_input(f"\{dia\}", "08:00")\
        entradas.append(val)\
        \
    alm_var = st.text_input("Hor\'e1rio de Almo\'e7o (Padr\'e3o para a semana)", "12:00")\
    \
    for i, ent_dia in enumerate(entradas):\
        escala_semanal[i] = calcular_quarteto(ent_dia, alm_var)\
    \
    txt_horario_cabecalho = "Vari\'e1vel (7h)"\
\
st.markdown("---")\
st.subheader("\uc0\u55356 \u57302 \u65039  Feriados")\
feriados_str = st.text_input("Dias de feriado (separe por v\'edrgula, ex: 15, 25)", "")\
feriados = []\
if feriados_str:\
    try:\
        feriados = [int(x.strip()) for x in feriados_str.split(",") if x.strip().isdigit()]\
    except:\
        st.error("Erro ao ler feriados. Use apenas n\'fameros separados por v\'edrgula.")\
\
# --- BOT\'c3O DE GERAR ---\
if st.button("Gerar Documento", type="primary"):\
    try:\
        doc = Document(ARQUIVO_MODELO)\
        \
        # 1. Preenchimento de Cabe\'e7alho\
        preencher_campo_seguro(doc, "Funcion\'e1rio", nome_func)\
        preencher_campo_seguro(doc, "Fun\'e7\'e3o", cargo_func)\
        preencher_campo_seguro(doc, "Emiss\'e3o", data_emissao, tamanho_fonte=9)\
        preencher_campo_seguro(doc, "Hor\'e1rio", txt_horario_cabecalho, tamanho_fonte=9)\
        \
        # 2. Marcar M\'eas\
        mapa_meses = \{\
            1: ["dez", "jan"], 2: ["jan", "fev"], 3: ["fev", "mar"], 4: ["mar", "abr"],\
            5: ["abr", "mai"], 6: ["mai", "jun"], 7: ["jun", "jul"], 8: ["jul", "ago"],\
            9: ["ago", "set"], 10: ["set", "out"], 11: ["out", "nov"], 12: ["nov", "dez"]\
        \}\
        mapa_texto = \{\
             1: "Dez/Jan", 2: "Jan/Fev", 3: "Fev/Mar", 4: "Mar/Abr", 5: "Abr/Mai", 6: "Mai/Jun",\
             7: "Jun/Jul", 8: "Jul/Ago", 9: "Ago/Set", 10: "Set/Out", 11: "Out/Nov", 12: "Nov/Dez"\
        \}\
        \
        termos = mapa_meses[mes_final]\
        txt_mes = mapa_texto[mes_final]\
        \
        for table in iterar_todas_as_tabelas(doc):\
            for row in table.rows:\
                for cell in row.cells:\
                    if all(t in cell.text.lower() for t in termos):\
                        limpar_e_escrever(cell, f"\{CHECKBOX_MARCADO\} \{txt_mes\}", tamanho=9, alinhamento=0)\
        \
        # 3. Marcar Base\
        marcar_base_preservando_linhas(doc, texto_base)\
        \
        # 4. Preencher Tabela de Dias\
        if mes_final == 1: mi, ai = 12, ano_final - 1\
        else: mi, ai = mes_final - 1, ano_final\
        \
        C_DIA, C_INI, C_ALM_IDA, C_ALM_VOL, C_FIM, C_ASS = 0, 2, 3, 4, 6, 7\
        \
        for table in doc.tables:\
            for row in table.rows:\
                try:\
                    txt_dia = row.cells[C_DIA].text.strip()\
                    if not txt_dia.isdigit(): continue\
                    dia = int(txt_dia)\
                    \
                    mc, ac = (mi, ai) if dia > 15 else (mes_final, ano_final)\
                    try: dt = date(ac, mc, dia)\
                    except: continue\
                    \
                    wd = dt.weekday()\
                    cells_alvo = [C_INI, C_ALM_IDA, C_ALM_VOL, C_FIM]\
                    if len(row.cells) <= C_FIM: cells_alvo = [1,2,3,4]\
\
                    # Prioridade: Feriado > Fim de Semana > Dia \'datil\
                    if dia in feriados:\
                        for i in cells_alvo: \
                            if i < len(row.cells): limpar_e_escrever(row.cells[i], "FERIADO")\
                        if C_ASS < len(row.cells): limpar_e_escrever(row.cells[C_ASS], "")\
                    elif wd == 5:\
                        for i in cells_alvo: \
                            if i < len(row.cells): limpar_e_escrever(row.cells[i], "S\'c1BADO")\
                        if C_ASS < len(row.cells): limpar_e_escrever(row.cells[C_ASS], "")\
                    elif wd == 6:\
                        for i in cells_alvo: \
                            if i < len(row.cells): limpar_e_escrever(row.cells[i], "DOMINGO")\
                        if C_ASS < len(row.cells): limpar_e_escrever(row.cells[C_ASS], "")\
                    else:\
                        horarios = escala_semanal.get(wd)\
                        if horarios:\
                            for k, idx in enumerate(cells_alvo):\
                                if idx < len(row.cells): limpar_e_escrever(row.cells[idx], horarios[k])\
                            if C_ASS < len(row.cells): inserir_assinatura(row.cells[C_ASS], ARQUIVO_ASSINATURA)\
                except IndexError: pass\
        \
        forcar_uma_pagina(doc)\
        \
        # SALVAR EM MEM\'d3RIA PARA DOWNLOAD\
        buffer = io.BytesIO()\
        doc.save(buffer)\
        buffer.seek(0)\
        \
        st.success("Documento gerado com sucesso!")\
        st.download_button(\
            label="\uc0\u11015 \u65039  Baixar Folha de Ponto",\
            data=buffer,\
            file_name=f"Folha_Ponto_\{mes_final\}_\{ano_final\}.docx",\
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"\
        )\
        \
    except Exception as e:\
        st.error(f"Erro ao gerar documento: \{e\}")}